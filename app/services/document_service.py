#!/usr/bin/env python3
"""
Document Service

Handles document parsing, rebuilding, and track changes generation.
Wraps functionality from parse_docx.py and rebuild_docx.py.
"""

import json
import shutil
import re
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

# Try importing redlines for track changes
try:
    from redlines import Redlines
    HAS_REDLINES = True
except ImportError:
    HAS_REDLINES = False


def get_paragraph_style_info(paragraph):
    """Extract style information from a paragraph."""
    style_name = paragraph.style.name if paragraph.style else "Normal"

    numbering_info = None
    if paragraph._p.pPr is not None:
        numPr = paragraph._p.pPr.find(qn('w:numPr'))
        if numPr is not None:
            ilvl = numPr.find(qn('w:ilvl'))
            numId = numPr.find(qn('w:numId'))
            if ilvl is not None and numId is not None:
                numbering_info = {
                    "level": int(ilvl.get(qn('w:val'))),
                    "numId": numId.get(qn('w:val'))
                }

    return {
        "style": style_name,
        "numbering": numbering_info,
        "is_heading": style_name.lower().startswith("heading")
    }


def extract_section_number(text):
    """Extract section number from paragraph text."""
    text = text.strip()

    patterns = [
        (r'^(ARTICLE\s+[IVXLCDM]+)[.\s:]+(.*)$', 'article'),
        (r'^(Article\s+[IVXLCDM]+)[.\s:]+(.*)$', 'article'),
        (r'^(ARTICLE\s+\d+)[.\s:]+(.*)$', 'article'),
        (r'^(Article\s+\d+)[.\s:]+(.*)$', 'article'),
        (r'^(Section\s+[\d]+\.[\d\.A-Za-z\(\)]+)[.\s:]+(.*)$', 'section'),
        (r'^(Section\s+[\d]+)[.\s:]+(.*)$', 'section'),
        (r'^(SECTION\s+[\d]+\.[\d\.A-Za-z\(\)]+)[.\s:]+(.*)$', 'section'),
        (r'^(SECTION\s+[\d]+)[.\s:]+(.*)$', 'section'),
        (r'^(\d+\.\d+\.\d+\.?\s*)(.*)$', 'subsub'),
        (r'^(\d+\.\d+\.?\s*)(.*)$', 'sub'),
        (r'^(\d+\.)\s+(.*)$', 'top'),
        (r'^([A-Z]\.)\s+(.*)$', 'letter_upper'),
        (r'^([a-z]\.)\s+(.*)$', 'letter_lower'),
        (r'^\(([A-Z])\)\s*(.*)$', 'paren_upper'),
        (r'^\(([a-z])\)\s*(.*)$', 'paren_lower'),
        (r'^\((\d+)\)\s*(.*)$', 'paren_num'),
        (r'^\(([ivxlcdm]+)\)\s*(.*)$', 'roman_lower'),
        (r'^\(([IVXLCDM]+)\)\s*(.*)$', 'roman_upper'),
    ]

    for pattern, num_type in patterns:
        match = re.match(pattern, text, re.IGNORECASE if num_type.startswith('article') or num_type.startswith('section') else 0)
        if match:
            section_num = match.group(1).strip()
            remaining = match.group(2).strip() if match.lastindex >= 2 else ""
            if num_type.startswith('paren') or num_type.startswith('roman'):
                if not section_num.startswith('('):
                    section_num = f"({section_num})"
            return (section_num, remaining, num_type)

    return (None, text, None)


def extract_caption(text, max_length=60):
    """Extract a caption from paragraph text."""
    section_num, remaining, _ = extract_section_number(text)
    text_to_use = remaining if remaining else text

    caption_match = re.match(r'^([^.]+\.)\s{2,}', text_to_use)
    if caption_match:
        return caption_match.group(1).strip()

    first_sentence = re.match(r'^([^.]+\.)', text_to_use)
    if first_sentence and len(first_sentence.group(1)) <= max_length:
        return first_sentence.group(1).strip()

    if len(text_to_use) > max_length:
        return text_to_use[:max_length].strip() + "..."

    return text_to_use.strip() if text_to_use.strip() else None


class SectionTracker:
    """Tracks the current section hierarchy as we parse the document."""

    def __init__(self):
        self.hierarchy = []
        self.counters = {}
        self.last_level = -1
        self.last_numId = None

    def update(self, numbering_level, section_num, caption, numId=None):
        """Update hierarchy based on new section encountered."""
        if numbering_level is not None and section_num is None:
            if numId != self.last_numId:
                self.counters = {}
                self.last_numId = numId

            levels_to_remove = [l for l in self.counters if l > numbering_level]
            for l in levels_to_remove:
                del self.counters[l]

            self.counters[numbering_level] = self.counters.get(numbering_level, 0) + 1
            section_num = self._generate_section_number(numbering_level)
            level = numbering_level

        elif section_num is not None:
            if numbering_level is not None:
                level = numbering_level
            else:
                if re.match(r'^(Article|ARTICLE|Section|SECTION)', section_num):
                    level = 0
                elif re.match(r'^\d+\.\d+\.\d+', section_num):
                    level = 2
                elif re.match(r'^\d+\.\d+', section_num):
                    level = 1
                elif re.match(r'^\d+\.', section_num):
                    level = 0
                elif re.match(r'^\([ivx]+\)', section_num, re.IGNORECASE):
                    level = 2
                elif re.match(r'^\([a-z]\)', section_num, re.IGNORECASE):
                    level = 2
                else:
                    level = self.last_level + 1 if self.last_level >= 0 else 0
        else:
            return

        self.hierarchy = self.hierarchy[:level]
        self.hierarchy.append({
            "level": level,
            "number": section_num,
            "caption": caption
        })
        self.last_level = level

    def _generate_section_number(self, level):
        """Generate section number string based on counters."""
        if level == 0:
            return f"{self.counters.get(0, 1)}."
        elif level == 1:
            count = self.counters.get(1, 1)
            letter = chr(ord('A') + count - 1) if count <= 26 else f"A{count-26}"
            return f"{letter}."
        elif level == 2:
            count = self.counters.get(2, 1)
            roman = self._to_roman(count).lower()
            return f"({roman})"
        else:
            count = self.counters.get(level, 1)
            return f"({count})"

    def _to_roman(self, num):
        """Convert integer to roman numeral."""
        val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
        syms = ['M', 'CM', 'D', 'CD', 'C', 'XC', 'L', 'XL', 'X', 'IX', 'V', 'IV', 'I']
        roman_num = ''
        for i in range(len(val)):
            while num >= val[i]:
                num -= val[i]
                roman_num += syms[i]
        return roman_num

    def get_current_hierarchy(self):
        """Return copy of current hierarchy."""
        return list(self.hierarchy)

    def get_section_ref(self):
        """Return concise section reference for manifest."""
        if not self.hierarchy:
            return None
        parts = [item["number"].rstrip('.') for item in self.hierarchy]
        return "".join(parts)


def extract_defined_terms(text):
    """Identify potential defined terms."""
    quoted_terms = re.findall(r'"([A-Z][^"]+)"', text)
    paren_terms = re.findall(r'\((?:the\s+)?"([A-Z][^"]+)"\)', text)
    return list(set(quoted_terms + paren_terms))


def iter_block_items(document):
    """Iterate through document body items in order."""
    parent = document.element.body
    for child in parent.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, document)
        elif child.tag == qn('w:tbl'):
            yield Table(child, document)


def process_table(table, start_id, section_tracker):
    """Process a table and return structured data."""
    table_data = {
        "type": "table",
        "id": f"tbl_{start_id}",
        "rows": [],
        "section_hierarchy": section_tracker.get_current_hierarchy()
    }

    para_id = start_id
    for row_idx, row in enumerate(table.rows):
        row_data = []
        for cell_idx, cell in enumerate(row.cells):
            cell_paragraphs = []
            for para in cell.paragraphs:
                para_id += 1
                text = para.text.strip()
                section_num, remaining, num_type = extract_section_number(text)
                caption = extract_caption(text)

                cell_paragraphs.append({
                    "id": f"p_{para_id}",
                    "text": text,
                    "section_number": section_num,
                    "caption": caption,
                    "style_info": get_paragraph_style_info(para),
                    "section_hierarchy": section_tracker.get_current_hierarchy()
                })
            row_data.append({
                "cell_id": f"cell_{row_idx}_{cell_idx}",
                "paragraphs": cell_paragraphs
            })
        table_data["rows"].append(row_data)

    return table_data, para_id


def parse_document(docx_path) -> Dict[str, Any]:
    """
    Parse a .docx file and extract structured content with section tracking.
    """
    docx_path = Path(docx_path)
    doc = Document(str(docx_path))
    section_tracker = SectionTracker()

    result = {
        "source_file": str(docx_path),
        "metadata": {"core_properties": {}},
        "content": [],
        "defined_terms": [],
        "sections": [],
        "exhibits": []
    }

    # Extract core properties
    try:
        props = doc.core_properties
        result["metadata"]["core_properties"] = {
            "title": props.title or "",
            "author": props.author or "",
            "created": str(props.created) if props.created else "",
            "modified": str(props.modified) if props.modified else ""
        }
    except Exception:
        pass

    para_id = 0
    all_defined_terms = set()

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            para_id += 1
            para_text = block.text.strip()
            style_info = get_paragraph_style_info(block)

            section_num, remaining, num_type = extract_section_number(para_text)
            caption = extract_caption(para_text)

            numbering_level = style_info["numbering"]["level"] if style_info["numbering"] else None
            numId = style_info["numbering"]["numId"] if style_info["numbering"] else None

            if numbering_level is not None or section_num or style_info["is_heading"]:
                section_tracker.update(numbering_level, section_num, caption, numId)

            para_data = {
                "type": "paragraph",
                "id": f"p_{para_id}",
                "text": para_text,
                "section_number": section_num,
                "section_ref": section_tracker.get_section_ref(),
                "caption": caption,
                "style_info": style_info,
                "section_hierarchy": section_tracker.get_current_hierarchy()
            }

            if style_info["is_heading"] or (section_num and num_type in ['article', 'section', 'top']):
                result["sections"].append({
                    "id": f"sec_{para_id}",
                    "number": section_num,
                    "title": caption or para_text[:50],
                    "para_id": f"p_{para_id}",
                    "hierarchy": section_tracker.get_current_hierarchy()
                })

            if re.match(r'^EXHIBIT\s+[A-Z0-9]', para_text, re.IGNORECASE):
                result["exhibits"].append({
                    "id": f"ex_{para_id}",
                    "title": para_text,
                    "start_para_id": f"p_{para_id}"
                })

            terms = extract_defined_terms(para_text)
            all_defined_terms.update(terms)

            result["content"].append(para_data)

        elif isinstance(block, Table):
            table_data, para_id = process_table(block, para_id, section_tracker)
            result["content"].append(table_data)

    result["defined_terms"] = sorted(list(all_defined_terms))

    return result


def replace_paragraph_text(paragraph, new_text):
    """Replace paragraph text while preserving formatting."""
    if not paragraph.runs:
        paragraph.text = new_text
        return

    first_run = paragraph.runs[0]
    first_run_format = {
        'bold': first_run.bold,
        'italic': first_run.italic,
        'underline': first_run.underline,
        'font_name': first_run.font.name,
        'font_size': first_run.font.size,
    }

    for run in paragraph.runs:
        run.text = ""

    first_run.text = new_text

    first_run.bold = first_run_format['bold']
    first_run.italic = first_run_format['italic']
    first_run.underline = first_run_format['underline']
    if first_run_format['font_name']:
        first_run.font.name = first_run_format['font_name']
    if first_run_format['font_size']:
        first_run.font.size = first_run_format['font_size']


def rebuild_document(original_path, revisions: Dict[str, Dict], output_path) -> int:
    """
    Rebuild document with revisions applied.

    Args:
        original_path: Path to original .docx file
        revisions: Dict mapping para_id to revision data
        output_path: Path for output .docx file

    Returns:
        Number of changes made
    """
    # Build lookup of revised content
    revised_lookup = {}
    for para_id, revision in revisions.items():
        if revision.get('accepted', False):
            revised_lookup[para_id] = revision.get('revised', '')

    # Copy original document
    shutil.copy2(original_path, output_path)

    # Open the copy and apply changes
    doc = Document(str(output_path))

    para_id = 0
    changes_made = 0

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            para_id += 1
            para_key = f"p_{para_id}"

            if para_key in revised_lookup:
                original_text = block.text.strip()
                revised_text = revised_lookup[para_key]

                if original_text != revised_text:
                    replace_paragraph_text(block, revised_text)
                    changes_made += 1

        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para_id += 1
                        para_key = f"p_{para_id}"

                        if para_key in revised_lookup:
                            original_text = para.text.strip()
                            revised_text = revised_lookup[para_key]

                            if original_text != revised_text:
                                replace_paragraph_text(para, revised_text)
                                changes_made += 1

    doc.save(str(output_path))
    return changes_made


def generate_track_changes_docx(original_path, revisions: Dict[str, Dict], output_path) -> str:
    """
    Generate Word document with track changes using Python-Redlines.

    If redlines is not available, falls back to standard rebuild.
    """
    if HAS_REDLINES:
        try:
            # Build list of changes
            changes = []
            for para_id, revision in revisions.items():
                if revision.get('accepted', False):
                    changes.append({
                        'original': revision.get('original', ''),
                        'revised': revision.get('revised', '')
                    })

            # Use redlines to generate track changes
            # Note: This is a simplified implementation
            # Full implementation would need to map changes back to specific paragraphs
            return rebuild_document(original_path, revisions, output_path)
        except Exception:
            pass

    # Fallback to standard rebuild
    return rebuild_document(original_path, revisions, output_path)


def generate_inline_diff_html(original: str, revised: str) -> str:
    """Generate inline HTML diff for display."""
    try:
        import diff_match_patch as dmp_module
        dmp = dmp_module.diff_match_patch()
        diffs = dmp.diff_main(original, revised)
        dmp.diff_cleanupSemantic(diffs)

        html_pieces = []
        for op, data in diffs:
            clean_data = data.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br>')
            if op == 1:
                html_pieces.append(f'<ins class="diff-ins">{clean_data}</ins>')
            elif op == -1:
                html_pieces.append(f'<del class="diff-del">{clean_data}</del>')
            else:
                html_pieces.append(clean_data)
        return "".join(html_pieces)
    except ImportError:
        # Fallback to simple side-by-side if diff_match_patch not available
        return f'<del class="diff-del">{original}</del><ins class="diff-ins">{revised}</ins>'


def generate_manifest(revisions: Dict[str, Dict], representation: str, deal_context: str) -> str:
    """Generate markdown manifest of all changes."""
    lines = [
        "# Redline Manifest",
        "",
        f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"**Representation:** {representation}",
        "",
        f"**Deal Context:** {deal_context}" if deal_context else "",
        "",
        "---",
        "",
        f"## Summary",
        "",
        f"Total revisions: {len([r for r in revisions.values() if r.get('accepted')])}",
        "",
        "---",
        "",
        "## Changes",
        ""
    ]

    for para_id, revision in sorted(revisions.items()):
        if not revision.get('accepted'):
            continue

        lines.extend([
            f"### {para_id}",
            "",
            f"**Original:**",
            f"> {revision.get('original', '')[:200]}...",
            "",
            f"**Revised:**",
            f"> {revision.get('revised', '')[:200]}...",
            "",
            f"**Rationale:** {revision.get('rationale', 'N/A')}",
            "",
            "---",
            ""
        ])

    return "\n".join(lines)


def generate_transmittal(revisions: Dict, flags: List, representation: str, deal_context: str) -> str:
    """Generate transmittal email text."""
    accepted_count = len([r for r in revisions.values() if r.get('accepted')])

    lines = [
        "DRAFT TRANSMITTAL EMAIL",
        "=" * 50,
        "",
        "Subject: Redlined Contract - [INSERT DEAL NAME]",
        "",
        "Dear [Client],",
        "",
        f"Please find attached our redlined version of the [Contract Type]. As {representation}'s counsel, "
        f"we have made {accepted_count} revisions to protect your interests.",
        "",
        "KEY CHANGES:",
        ""
    ]

    # Group revisions by type/category if available
    change_types = {}
    for para_id, rev in revisions.items():
        if rev.get('accepted'):
            rationale = rev.get('rationale', 'General revision')
            if rationale not in change_types:
                change_types[rationale] = 0
            change_types[rationale] += 1

    for rationale, count in sorted(change_types.items(), key=lambda x: -x[1])[:5]:
        lines.append(f"- {rationale} ({count} instances)")

    lines.extend([
        "",
        "ITEMS FLAGGED FOR YOUR REVIEW:",
        ""
    ])

    if flags:
        for flag in flags:
            lines.append(f"- Section {flag.get('section_ref', 'N/A')}: {flag.get('note', 'Please review')}")
    else:
        lines.append("- No items specifically flagged.")

    lines.extend([
        "",
        "Please let us know if you have any questions or would like to discuss any of these changes.",
        "",
        "Best regards,",
        "[Attorney Name]"
    ])

    return "\n".join(lines)


def generate_final_output(session_id: str, original_path: str, parsed_doc: Dict,
                          revisions: Dict, flags: List, representation: str,
                          deal_context: str) -> Dict[str, str]:
    """
    Generate all final output files.

    Returns dict with paths to generated files.
    """
    output_dir = Path(original_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)

    # Generate revised Word document
    docx_path = output_dir / 'revised.docx'
    changes_made = generate_track_changes_docx(original_path, revisions, docx_path)

    # Generate manifest
    manifest_path = output_dir / 'manifest.md'
    manifest_content = generate_manifest(revisions, representation, deal_context)
    with open(manifest_path, 'w', encoding='utf-8') as f:
        f.write(manifest_content)

    # Generate transmittal
    transmittal_path = output_dir / 'transmittal.txt'
    transmittal_content = generate_transmittal(revisions, flags, representation, deal_context)
    with open(transmittal_path, 'w', encoding='utf-8') as f:
        f.write(transmittal_content)

    return {
        'docx_path': str(docx_path),
        'manifest_path': str(manifest_path),
        'transmittal_path': str(transmittal_path),
        'changes_made': changes_made,
        'accepted_revisions': len([r for r in revisions.values() if r.get('accepted')]),
        'flags_count': len(flags)
    }


def generate_final_documents(session_id: str, original_path: str, parsed_doc: Dict,
                             revisions: Dict, author_name: str = "Contract Review Tool") -> Dict[str, Any]:
    """
    Generate final Word documents for export.

    Creates two documents:
    1. Track changes document showing all accepted revisions with markup
    2. Clean document with final text only (no markup)

    Args:
        session_id: Session identifier
        original_path: Path to original .docx file
        parsed_doc: Parsed document structure
        revisions: Dict mapping para_id to revision data
        author_name: Author name for track changes attribution

    Returns:
        Dict with paths to generated files and revision details
    """
    from lxml import etree
    from docx.oxml.ns import nsmap
    from copy import deepcopy

    output_dir = Path(original_path).parent
    output_dir.mkdir(parents=True, exist_ok=True)

    # Build lookup of accepted revisions
    accepted_revisions = {}
    revision_details = []

    for para_id, revision in revisions.items():
        if revision.get('accepted', False):
            accepted_revisions[para_id] = revision
            # Find section reference from parsed_doc
            section_ref = None
            para_text_preview = revision.get('original', '')[:100]
            for item in parsed_doc.get('content', []):
                if item.get('id') == para_id:
                    section_ref = item.get('section_ref', '')
                    para_text_preview = item.get('text', '')[:100]
                    break

            revision_details.append({
                'para_id': para_id,
                'section_ref': section_ref or para_id,
                'original_preview': para_text_preview + ('...' if len(revision.get('original', '')) > 100 else ''),
                'rationale': revision.get('rationale', 'No rationale provided')
            })

    # Generate track changes document using redlines
    track_changes_path = output_dir / f'{session_id}_track_changes.docx'
    clean_path = output_dir / f'{session_id}_clean.docx'

    if HAS_REDLINES and accepted_revisions:
        try:
            # Create track changes document
            _generate_track_changes_with_redlines(
                original_path,
                accepted_revisions,
                track_changes_path,
                author_name
            )
        except Exception as e:
            # Fallback to simple rebuild if redlines fails
            rebuild_document(original_path, revisions, track_changes_path)
    else:
        # Fallback to simple rebuild
        rebuild_document(original_path, revisions, track_changes_path)

    # Generate clean document (final text only)
    rebuild_document(original_path, revisions, clean_path)

    return {
        'track_changes_path': str(track_changes_path),
        'clean_path': str(clean_path),
        'revision_count': len(accepted_revisions),
        'revision_details': revision_details
    }


def _generate_track_changes_with_redlines(original_path: str, accepted_revisions: Dict,
                                          output_path: str, author_name: str) -> None:
    """
    Generate a Word document with track changes using python-redlines.

    This creates proper Word track changes that display correctly in Microsoft Word.
    """
    from redlines import Redlines

    # Copy original document first
    shutil.copy2(original_path, output_path)

    # Open the document
    doc = Document(str(output_path))

    # Track paragraph index
    para_id = 0

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            para_id += 1
            para_key = f"p_{para_id}"

            if para_key in accepted_revisions:
                revision = accepted_revisions[para_key]
                original_text = revision.get('original', '')
                revised_text = revision.get('revised', '')

                if original_text != revised_text:
                    # Use redlines to generate the diff
                    redline = Redlines(original_text, revised_text)

                    # Get the redline output as markdown-style text
                    # We need to apply this as Word track changes
                    _apply_track_changes_to_paragraph(block, original_text, revised_text, author_name)

        elif isinstance(block, Table):
            for row in block.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para_id += 1
                        para_key = f"p_{para_id}"

                        if para_key in accepted_revisions:
                            revision = accepted_revisions[para_key]
                            original_text = revision.get('original', '')
                            revised_text = revision.get('revised', '')

                            if original_text != revised_text:
                                _apply_track_changes_to_paragraph(para, original_text, revised_text, author_name)

    doc.save(str(output_path))


def _apply_track_changes_to_paragraph(paragraph, original_text: str, revised_text: str, author_name: str) -> None:
    """
    Apply track changes to a single paragraph using Word's revision markup.

    Creates proper <w:del> and <w:ins> elements for Word track changes.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from datetime import datetime
    import diff_match_patch as dmp_module

    # Use diff_match_patch to compute changes
    dmp = dmp_module.diff_match_patch()
    diffs = dmp.diff_main(original_text, revised_text)
    dmp.diff_cleanupSemantic(diffs)

    # Get first run's formatting if available
    first_run_format = None
    if paragraph.runs:
        first_run = paragraph.runs[0]
        first_run_format = {
            'bold': first_run.bold,
            'italic': first_run.italic,
            'underline': first_run.underline,
            'font_name': first_run.font.name,
            'font_size': first_run.font.size,
        }

    # Clear existing paragraph content
    p = paragraph._p
    for child in list(p):
        if child.tag == qn('w:r'):
            p.remove(child)

    # Current timestamp for revisions
    rev_date = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

    # Build new content with track changes
    for op, text in diffs:
        if op == 0:  # Equal - no change
            run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            _apply_run_formatting(rPr, first_run_format)
            run.append(rPr)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = text
            run.append(t)
            p.append(run)

        elif op == -1:  # Deletion
            del_elem = OxmlElement('w:del')
            del_elem.set(qn('w:id'), str(hash(text) % 100000))
            del_elem.set(qn('w:author'), author_name)
            del_elem.set(qn('w:date'), rev_date)

            run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            _apply_run_formatting(rPr, first_run_format)
            run.append(rPr)

            delText = OxmlElement('w:delText')
            delText.set(qn('xml:space'), 'preserve')
            delText.text = text
            run.append(delText)

            del_elem.append(run)
            p.append(del_elem)

        elif op == 1:  # Insertion
            ins_elem = OxmlElement('w:ins')
            ins_elem.set(qn('w:id'), str(hash(text) % 100000))
            ins_elem.set(qn('w:author'), author_name)
            ins_elem.set(qn('w:date'), rev_date)

            run = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            _apply_run_formatting(rPr, first_run_format)
            run.append(rPr)

            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = text
            run.append(t)

            ins_elem.append(run)
            p.append(ins_elem)


def _apply_run_formatting(rPr, format_dict):
    """Apply formatting properties to a run properties element."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if not format_dict:
        return

    if format_dict.get('bold'):
        b = OxmlElement('w:b')
        rPr.append(b)

    if format_dict.get('italic'):
        i = OxmlElement('w:i')
        rPr.append(i)

    if format_dict.get('underline'):
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    if format_dict.get('font_name'):
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), format_dict['font_name'])
        rFonts.set(qn('w:hAnsi'), format_dict['font_name'])
        rPr.append(rFonts)

    if format_dict.get('font_size'):
        sz = OxmlElement('w:sz')
        # Font size in half-points
        sz.set(qn('w:val'), str(int(format_dict['font_size'].pt * 2)))
        rPr.append(sz)
